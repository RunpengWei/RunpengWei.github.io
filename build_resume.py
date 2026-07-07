from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PHOTO = '/Users/kanyun/Desktop/resume-site/photo.jpg'
OUTPUT = '/Users/kanyun/Desktop/resume-site/CV_WEI Runpeng_家辉教育语文教师.docx'
WEBSITE = 'https://runpengwei.github.io/'
CPL = 48

# Typography palette
ACCENT = RGBColor(0x1A, 0x47, 0x7A)
TEXT = RGBColor(0x33, 0x33, 0x33)
MUTED = RGBColor(0x66, 0x66, 0x66)
LIGHT = RGBColor(0x88, 0x88, 0x88)

FONT_BODY = '宋体'
FONT_HEADING = '黑体'

SZ_NAME = 20
SZ_CONTACT = 9
SZ_SECTION = 11
SZ_ENTRY = 10
SZ_DATE = 9
SZ_BODY = 9

LINE_SPACING = 1.15
CONTENT_WIDTH = Cm(17.0)  # tab stop for right-aligned dates


def add_hyperlink(paragraph, text, url, size=SZ_CONTACT):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '1A477A')
    rPr.append(c)
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), FONT_BODY)
    rf.set(qn('w:eastAsia'), FONT_BODY)
    rPr.append(rf)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size * 2)))
    rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_ea(run, name):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def style_run(run, size=SZ_BODY, bold=False, color=TEXT, font=FONT_BODY):
    set_ea(run, font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def set_para_format(p, before=0, after=0, align=None, justify=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SPACING
    if align is not None:
        pf.alignment = align
    elif justify:
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def set_cell_valign(cell, align='top'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def set_table_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


# Content blocks (unchanged)
B = {
    'edu': '绩点 3.71/4.00（本）、3.91/4.00（硕）；核心课程涵盖中国古代史、中国佛教经典选读、图像阅读、考古学通论等，系统训练文献研读与学术写作，为语文教学文本细读奠定人文根基；高考语文 136 分，持有高中语文教师资格证，雅思 7.5，燕京学堂全英文项目在读，具备跨文化阅读视野。',
    'yb': '参与儿童音频内容生产全流程：依据儿童认知特点将经典名著、古诗词转化为"可听、可懂、可回味"的讲解文本，负责审核与上线，在准确性与童趣表达间反复打磨；坚持"适切性+挑战性"分层设计，低年级重故事韵律，高年级引入背景知识与思辨追问；开展儿童有声书市场调研，梳理竞品结构与用户互动模式，认识到学生习惯碎片化接收信息，语文教育需探索"深度内容+轻量载体"融合路径，积累可感知体验的一线教研经验',
    'pup': '普林斯顿大学出版社（2026.01-2026.03）：独立运营小红书官方账号，以书单策划、单本导读、热点书评进行人文内容传播，周更 2-3 条，单月曝光量增长 66%；撰写推文与视频脚本，将学术出版物转化为年轻读者可读表达，锤炼把经典与知识转化为学生愿意接近、能够消化的语言之"翻译"能力',
    'tutor1': '累计服务 20 名学生，高峰期同时辅导 8 人；入学前进行学情诊断（阅读速度、写作结构、应试策略、学习动机），为每位学生定制阶段性课程方案，薄弱学生以精读仿写夯实语感，优等生引入专题阅读与论述训练，成绩普遍提升，排名最高跃升 60%，多名学生反馈愿意主动阅读，在提分与素养培育间寻求平衡。',
    'tutor2': '坚持"分数是结果，素养是根基"，不牺牲阅读广度换取短期技巧，侧重阅读写作与文学素养培养，帮助学生建立可持续学习方法与表达自信，多名学生反馈"愿意主动拿起书"，在提分与人文积淀之间寻求平衡，见效',
    'zx': '统筹蓝信封书信与寒暑假支教，联络 70 余名志愿者与太平中学学生建立书信往来，理解学生需要被倾听、被看见的安全感；赴河北某中学支教，设计"基础巩固+兴趣拓展"双线方案，将古诗文与生活场景、地方文化联结，参与备课授课与学情反馈；口述史调研整理 10 万余字资料，获第 31 届挑战杯二等奖。',
    'acad': '参与法藏敦煌文献整理编目，小组已梳理逾 4000 件资料、出版图录 10 余册，严谨求证与文本细节敏感度可迁移至篇章解读；论文《呦呦鹿鸣》获学术文化节三等奖。克孜尔石窟项目：撰写叙事脚本与展览信息卡，将壁画考古成果编织为故事线，实践"让学术走向公众"的内容转化，与语文沉浸式阅读教学相通。',
    'honor': '北京大学优秀毕业生、优秀学生干部、十佳志愿者标兵｜第 31 届挑战杯二等奖"支教团口述史研究"｜学术文化节三等奖《呦呦鹿鸣：从印度到中国的鹿形象》｜课程优秀论文奖《冥府絮语——〈吉尔伽美什史诗〉再创作》，综合体现扎实学业功底、科研素养与人文综合能力与全面发展的综合素养与特长水平',
    'self': '文史功底扎实，高考语文 136 分，持有高中语文教资，擅长文本细读、写作指导与文学赏析；四年家教、支教与教研经历，熟悉初高中认知特点，善于学情诊断与个性化方案设计；我认为新教培时代，学生最需要的不仅是分数，更是在 AI 与信息过载背景下培养思辨力与表达自信，让语文课堂连接经典与生活、知识',
}


def report_lengths():
    for key, text in B.items():
        lines = (len(text) + CPL - 1) // CPL
        waste = lines * CPL - len(text)
        print(f'{key}: {len(text)}c {lines}L waste={waste}')


def section(doc, title):
    p = doc.add_paragraph()
    set_para_format(p, before=5, after=1)
    style_run(p.add_run(title), SZ_SECTION, True, ACCENT, FONT_HEADING)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), '4')
    b.set(qn('w:space'), '2')
    b.set(qn('w:color'), '1A477A')
    pBdr.append(b)
    pPr.append(pBdr)


def entry(doc, left, right):
    p = doc.add_paragraph()
    set_para_format(p, before=2, after=0)
    style_run(p.add_run(left), SZ_ENTRY, True)
    p.add_run('\t')
    style_run(p.add_run(right), SZ_DATE, False, MUTED)
    p.paragraph_format.tab_stops.add_tab_stop(CONTENT_WIDTH, WD_TAB_ALIGNMENT.RIGHT)


def bullet(doc, text):
    p = doc.add_paragraph()
    set_para_format(p, before=0, after=0, justify=True)
    pf = p.paragraph_format
    pf.left_indent = Cm(0.38)
    pf.first_line_indent = Cm(-0.38)
    style_run(p.add_run('• ' + text), SZ_BODY)


def plain(doc, text):
    p = doc.add_paragraph()
    set_para_format(p, before=1, after=0, justify=True)
    style_run(p.add_run(text), SZ_BODY)


def build_header(doc):
    table = doc.add_table(rows=1, cols=2)
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    tblPr.append(borders)

    set_table_col_widths(table, [Cm(13.8), Cm(3.2)])

    lc, rc = table.rows[0].cells[0], table.rows[0].cells[1]
    set_cell_valign(lc, 'top')
    set_cell_valign(rc, 'top')

    # Name
    pn = lc.paragraphs[0]
    set_para_format(pn, before=0, after=2)
    style_run(pn.add_run('韦润芃'), SZ_NAME, True, ACCENT, FONT_HEADING)

    # Phone + email
    pc = lc.add_paragraph()
    set_para_format(pc, before=0, after=1)
    style_run(pc.add_run('手机/微信：'), SZ_CONTACT, False, LIGHT)
    style_run(pc.add_run('(86) 185-2259-1129'), SZ_CONTACT, False, TEXT)
    style_run(pc.add_run('    邮箱：'), SZ_CONTACT, False, LIGHT)
    style_run(pc.add_run('2501213758@stu.pku.edu.cn'), SZ_CONTACT, False, TEXT)

    # Website
    pw = lc.add_paragraph()
    set_para_format(pw, before=0, after=3)
    style_run(pw.add_run('个人简历网站：'), SZ_CONTACT, False, LIGHT)
    add_hyperlink(pw, WEBSITE, WEBSITE, SZ_CONTACT)

    # Photo — top-right, standard 一寸照比例
    pp = rc.paragraphs[0]
    pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_format(pp)
    pp.add_run().add_picture(PHOTO, width=Cm(2.6), height=Cm(3.64))


def build():
    report_lengths()

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.0)
    sec.bottom_margin = Cm(0.9)
    sec.left_margin = Cm(1.6)
    sec.right_margin = Cm(1.6)

    # Default document style
    normal = doc.styles['Normal']
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    normal.font.size = Pt(SZ_BODY)
    normal.font.color.rgb = TEXT

    build_header(doc)

    section(doc, '教育背景')
    entry(doc, '北京大学  本科·历史学系（"敦煌与中外文明"项目）', '2021.09 - 2025.06')
    entry(doc, '北京大学  硕士·燕京学堂·中国学（历史与考古方向）', '2025.09 - 2027.06')
    bullet(doc, B['edu'])

    section(doc, '实习经历')
    entry(doc, '猿辅导·斑马App  语文教研实习生', '2026.04 - 至今')
    bullet(doc, B['yb'])
    bullet(doc, B['pup'])

    section(doc, '教学与实践经历')
    entry(doc, '语文家教辅导  初中至高中一对一', '2021.09 - 至今')
    bullet(doc, B['tutor1'])
    bullet(doc, B['tutor2'])
    entry(doc, '北京大学教育知行社  项目部副部长', '2022.09 - 2023.04')
    bullet(doc, B['zx'])

    section(doc, '学术研究')
    entry(doc, '法藏敦煌文献整理小组  成员（荣新江老师指导）', '2024.02 - 至今')
    bullet(doc, B['acad'])

    section(doc, '荣誉奖励')
    plain(doc, B['honor'])

    section(doc, '自我评价')
    bullet(doc, B['self'])

    doc.save(OUTPUT)

    total = 6
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        total += (len(t) + CPL - 1) // CPL
    print('Saved:', OUTPUT)
    print('TOTAL visual lines:', total)


if __name__ == '__main__':
    build()
